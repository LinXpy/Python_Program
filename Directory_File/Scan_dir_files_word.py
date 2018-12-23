
# Backup the files in the given path ####


import os
from time import strftime, localtime
import docx


fsizedicr = {"B": 1,                        # B=Byte
             "KB": float(1)/1024,
             "MB": float(1)/(1024*1024),
             "GB": float(1)/(1024*1024*1024)
             }

scan_path = input("Please input the directory path you want to scan(like D:\somedir\somedir): ")
save_path = input("Please input the absolute-path filename you want the result file to be saved(like D:\somedir\somedir\filename.docx): ")
document = docx.Document()


for dir_path, dir_names, file_names in os.walk(scan_path):

    document.add_heading("Files in directory: " + dir_path, level=1)
    document.add_paragraph('='*60)
    document.add_paragraph("Files and Folders in The Current Directory: ")
#    document.add_paragraph(os.listdir(dir_path))   # no space between the elements
    dirlists = ''
    for elem in os.listdir(dir_path):
        dirlists += elem + ' ;  '
    document.add_paragraph(dirlists)

    for file in file_names:
        file_abpath = os.path.join(dir_path, file)
        file_atime = os.path.getatime(file_abpath)  # get last access time
        file_mtime = os.path.getmtime(file_abpath)  # get last modification time
        file_ctime = os.path.getctime(file_abpath)  # get creation time for Windows, last metadata change time for Linux
        file_size = os.path.getsize(file_abpath)    # get size
        access_time = strftime("%Y/%m/%d %H:%M:%S", localtime(file_atime))
        modification_time = strftime("%Y/%m/%d %H:%M:%S", localtime(file_mtime))
        creation_time = strftime("%Y/%m/%d %H:%M:%S", localtime(file_ctime))

        if file_size > 1024*1024*1024:
            sfile = str(round(fsizedicr['GB']*file_size, 2))+'GB'
        elif file_size > 1024*1024:
            sfile = str(round(fsizedicr['MB']*file_size, 2))+'MB'
        elif file_size > 1024:
            sfile = str(round(fsizedicr['KB']*file_size, 2))+'KB'
        else:
            sfile = str(round(fsizedicr['B']*file_size, 2))+'B'

        file_info = "{:<30}".format(file) + "{:<10}".format(sfile) + "{:<}".format(creation_time)    # combine file information string
        document.add_paragraph(file_info)

    document.add_paragraph('=' * 60)

document.add_page_break()
# document.save(r"D:\\Lang_Prog\\Python_Program\\file_backup.docx")
document.save(save_path)




